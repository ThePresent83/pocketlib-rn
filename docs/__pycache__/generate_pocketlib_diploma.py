from pathlib import Path
from textwrap import wrap

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Mm, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(__file__).resolve().parent
OUTPUT = ROOT / "PocketLib_Дипломная_документация.docx"
ARCHITECTURE_IMAGE = ROOT / "pocketlib_architecture.png"
DATABASE_IMAGE = ROOT / "pocketlib_database.png"
ROLES_IMAGE = ROOT / "pocketlib_roles.png"


def font(size=28, bold=False):
    name = "arialbd.ttf" if bold else "arial.ttf"
    path = Path("C:/Windows/Fonts") / name
    return ImageFont.truetype(str(path), size=size)


def draw_box(draw, xy, title, lines, fill, outline="#24445f"):
    x1, y1, x2, y2 = xy
    draw.rounded_rectangle(xy, radius=18, fill=fill, outline=outline, width=3)
    draw.text((x1 + 18, y1 + 14), title, font=font(24, True), fill="#153047")
    y = y1 + 54
    for line in lines:
        for part in wrap(line, width=max(18, int((x2 - x1) / 14))):
            draw.text((x1 + 18, y), part, font=font(19), fill="#243746")
            y += 25


def arrow(draw, start, end, label=""):
    draw.line([start, end], fill="#2f6f8f", width=4)
    x2, y2 = end
    x1, y1 = start
    dx, dy = x2 - x1, y2 - y1
    length = max((dx * dx + dy * dy) ** 0.5, 1)
    ux, uy = dx / length, dy / length
    px, py = -uy, ux
    points = [
        (x2, y2),
        (x2 - ux * 18 + px * 9, y2 - uy * 18 + py * 9),
        (x2 - ux * 18 - px * 9, y2 - uy * 18 - py * 9),
    ]
    draw.polygon(points, fill="#2f6f8f")
    if label:
        mx, my = (x1 + x2) / 2, (y1 + y2) / 2
        draw.rounded_rectangle((mx - 70, my - 18, mx + 70, my + 16), radius=8, fill="#ffffff")
        draw.text((mx - 61, my - 13), label, font=font(16), fill="#2f6f8f")


def build_diagrams():
    image = Image.new("RGB", (1500, 930), "#f7fbfd")
    draw = ImageDraw.Draw(image)
    draw.text((40, 28), "Архитектура PocketLib", font=font(38, True), fill="#14344b")
    draw_box(draw, (70, 150, 420, 390), "Мобильное приложение", [
        "React Native + Expo",
        "Expo Router",
        "React Native Paper",
        "Встроенный ридер",
    ], "#dff1f8")
    draw_box(draw, (575, 110, 930, 355), "Локальные данные", [
        "SQLite: каталог и пользователи",
        "AsyncStorage: прогресс и закладки",
        "FileSystem: загруженные файлы",
    ], "#e9f6e5")
    draw_box(draw, (1080, 110, 1430, 355), "Project Gutenberg", [
        "Обложки",
        "Легальные открытые тексты",
        "Gutendex: поиск метаданных",
    ], "#fff2d9")
    draw_box(draw, (575, 555, 930, 810), "BK Library API", [
        "Node.js HTTP API",
        "Каталог XML KostanaySoft",
        "Фильтр официальных ссылок",
        "Кэш источников",
    ], "#f2e7fa")
    draw_box(draw, (1080, 555, 1430, 810), "Официальные источники", [
        "gov.kz",
        "data.egov.kz",
        "Сайты издательств",
        "Только разрешенные ссылки",
    ], "#fbe7ea")
    arrow(draw, (420, 250), (575, 230), "SQLite")
    arrow(draw, (930, 230), (1080, 230), "HTTPS")
    arrow(draw, (420, 330), (575, 650), "HTTP :3047")
    arrow(draw, (930, 680), (1080, 680), "HTTPS")
    image.save(ARCHITECTURE_IMAGE)

    image = Image.new("RGB", (1500, 1060), "#fbfcfe")
    draw = ImageDraw.Draw(image)
    draw.text((40, 28), "Логическая схема локальной базы данных", font=font(38, True), fill="#14344b")
    boxes = {
        "users": ((70, 130, 420, 350), ["id", "full_name", "email", "password", "role", "speciality_id", "course_id"]),
        "books": ((565, 130, 970, 490), ["id", "title", "author", "description", "cover_url", "file_path", "source", "gutenberg_id", "discipline_id", "course_id"]),
        "disciplines": ((1080, 130, 1430, 315), ["id", "name", "color"]),
        "courses": ((1080, 425, 1430, 625), ["id", "name", "year", "discipline_id"]),
        "specialities": ((70, 480, 420, 655), ["id", "name"]),
        "categories": ((70, 750, 420, 925), ["id", "name"]),
        "reading_history": ((565, 690, 970, 925), ["id", "user_id", "book_id", "last_opened", "progress"]),
        "search_cache": ((1080, 750, 1430, 925), ["query", "results_json", "timestamp"]),
    }
    colors = ["#dff1f8", "#e9f6e5", "#fff2d9", "#f2e7fa", "#fbe7ea", "#e9edf8", "#f5f0dd", "#e2f4f0"]
    for (title, (xy, fields)), fill in zip(boxes.items(), colors):
        draw_box(draw, xy, title, fields, fill)
    arrow(draw, (420, 260), (565, 250), "user")
    arrow(draw, (970, 245), (1080, 220), "discipline")
    arrow(draw, (970, 340), (1080, 515), "course")
    arrow(draw, (420, 575), (565, 210), "speciality")
    arrow(draw, (760, 490), (760, 690), "history")
    image.save(DATABASE_IMAGE)

    image = Image.new("RGB", (1500, 640), "#fbfcfe")
    draw = ImageDraw.Draw(image)
    draw.text((40, 28), "Ролевая модель PocketLib", font=font(38, True), fill="#14344b")
    draw_box(draw, (70, 145, 450, 535), "Студент", [
        "Просмотр библиотеки",
        "Поиск литературы",
        "Чтение текста",
        "Прогресс и закладки",
        "Настройка языка",
    ], "#dff1f8")
    draw_box(draw, (560, 145, 940, 535), "Преподаватель", [
        "Все функции студента",
        "Добавление материалов",
        "Назначение дисциплины",
        "Работа с локальными файлами",
    ], "#e9f6e5")
    draw_box(draw, (1050, 145, 1430, 535), "Администратор", [
        "Все функции преподавателя",
        "Управление пользователями",
        "Изменение ролей",
        "Удаление учетных записей",
        "Синхронизация каталогов",
    ], "#fff2d9")
    arrow(draw, (450, 340), (560, 340), "расширение")
    arrow(draw, (940, 340), (1050, 340), "расширение")
    image.save(ROLES_IMAGE)


def set_cell_text(cell, text, bold=False):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run(str(text))
    r.bold = bold
    r.font.name = "Times New Roman"
    r.font.size = Pt(11)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    repeat = OxmlElement("w:tblHeader")
    repeat.set(qn("w:val"), "true")
    tr_pr.append(repeat)


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    text = OxmlElement("w:instrText")
    text.set(qn("xml:space"), "preserve")
    text.text = " PAGE "
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, text, end])


def add_toc(paragraph):
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:instrText")
    text.set(qn("xml:space"), "preserve")
    text.text = ' TOC \\o "1-3" \\h \\z \\u '
    placeholder = OxmlElement("w:t")
    placeholder.text = "Оглавление обновится после открытия документа в Microsoft Word."
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, text, separate, placeholder, end])


def configure_document(doc):
    section = doc.sections[0]
    section.page_width = Mm(210)
    section.page_height = Mm(297)
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(3)
    section.right_margin = Cm(1.5)
    section.header_distance = Cm(1.2)
    section.footer_distance = Cm(1)
    section.different_first_page_header_footer = True
    add_page_number(section.footer.paragraphs[0])

    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    normal.font.size = Pt(14)
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    normal.paragraph_format.first_line_indent = Cm(1.25)
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    normal.paragraph_format.space_after = Pt(0)

    for name, size, bold in [("Title", 18, True), ("Heading 1", 16, True), ("Heading 2", 14, True), ("Heading 3", 14, True)]:
        style = doc.styles[name]
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        style.font.size = Pt(size)
        style.font.bold = bold
        style.paragraph_format.first_line_indent = Cm(0)
        style.paragraph_format.space_before = Pt(12)
        style.paragraph_format.space_after = Pt(6)
    doc.styles["Heading 1"].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.styles["Heading 2"].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    doc.styles["Heading 3"].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT

    code = doc.styles.add_style("Code", WD_STYLE_TYPE.PARAGRAPH)
    code.font.name = "Consolas"
    code._element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")
    code.font.size = Pt(9)
    code.paragraph_format.left_indent = Cm(0.7)
    code.paragraph_format.first_line_indent = Cm(0)
    code.paragraph_format.space_after = Pt(2)
    code.paragraph_format.line_spacing = 1

    settings = doc.settings.element
    update = OxmlElement("w:updateFields")
    update.set(qn("w:val"), "true")
    settings.append(update)


def add_paragraph(doc, text, style=None):
    p = doc.add_paragraph(style=style)
    p.add_run(text)
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.left_indent = Cm(0.75)
        p.paragraph_format.first_line_indent = Cm(0)
        p.add_run(item)


def add_numbered(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.left_indent = Cm(0.75)
        p.paragraph_format.first_line_indent = Cm(0)
        p.add_run(item)


def add_table(doc, headers, rows, widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for idx, header in enumerate(headers):
        set_cell_text(table.rows[0].cells[idx], header, True)
    set_repeat_table_header(table.rows[0])
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            set_cell_text(cells[idx], value)
    if widths:
        for row in table.rows:
            for idx, width in enumerate(widths):
                row.cells[idx].width = Cm(width)
    doc.add_paragraph()
    return table


def add_caption(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Cm(0)
    run = p.add_run(text)
    run.italic = True
    run.font.size = Pt(12)


def add_code(doc, code):
    for line in code.strip("\n").splitlines():
        p = doc.add_paragraph(style="Code")
        p.add_run(line)
    doc.add_paragraph()


def add_main_heading(doc, text):
    doc.add_page_break()
    doc.add_heading(text, level=1)


def title_page(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Cm(0)
    for text, size, bold in [
        ("МИНИСТЕРСТВО ПРОСВЕЩЕНИЯ РЕСПУБЛИКИ КАЗАХСТАН", 12, True),
        ("[НАИМЕНОВАНИЕ КОЛЛЕДЖА]", 12, True),
        ("[ОТДЕЛЕНИЕ / ЦИКЛОВАЯ КОМИССИЯ]", 12, False),
    ]:
        r = p.add_run(text + "\n")
        r.bold = bold
        r.font.size = Pt(size)

    for _ in range(5):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Cm(0)
    r = p.add_run("ДИПЛОМНЫЙ ПРОЕКТ\n")
    r.bold = True
    r.font.size = Pt(18)
    r = p.add_run("ПОЯСНИТЕЛЬНАЯ ЗАПИСКА\n\n")
    r.bold = True
    r.font.size = Pt(16)
    r = p.add_run("Разработка мобильного приложения электронной библиотеки\n«PocketLib»")
    r.bold = True
    r.font.size = Pt(16)

    for _ in range(5):
        doc.add_paragraph()
    table = doc.add_table(rows=4, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.RIGHT
    data = [
        ("Выполнил(а):", "[Ф.И.О. студента]"),
        ("Группа:", "[номер группы]"),
        ("Специальность:", "[наименование специальности]"),
        ("Руководитель:", "[Ф.И.О. руководителя]"),
    ]
    for row, values in zip(table.rows, data):
        set_cell_text(row.cells[0], values[0])
        set_cell_text(row.cells[1], values[1])
        row.cells[0].width = Cm(4)
        row.cells[1].width = Cm(7)
    for cell in table._element.xpath(".//w:tc"):
        tc_pr = cell.get_or_add_tcPr()
        borders = OxmlElement("w:tcBorders")
        for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
            tag = OxmlElement(f"w:{edge}")
            tag.set(qn("w:val"), "nil")
            borders.append(tag)
        tc_pr.append(borders)

    for _ in range(5):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Cm(0)
    p.add_run("[город] — 2026")
    doc.add_page_break()


def annotation(doc):
    doc.add_heading("АННОТАЦИЯ", level=1)
    add_paragraph(doc, "В дипломном проекте разработано мобильное приложение PocketLib, предназначенное для формирования электронной библиотеки колледжа, поиска открытой литературы, чтения текстовых материалов внутри приложения и управления локальным каталогом. Клиентская часть реализована на React Native и Expo с применением TypeScript. Для локального хранения данных используется SQLite, для прогресса чтения и закладок — AsyncStorage.")
    add_paragraph(doc, "Отдельный локальный сервис BK Library API создан на Node.js. Он импортирует метаданные из XML-выгрузки АБИС «Библиотечное дело» KostanaySoft, предоставляет методы просмотра каталога и фильтрует официальные ссылки на электронные учебники Казахстана. Полные тексты из XML не извлекаются, так как исходная выгрузка содержит преимущественно библиографические сведения.")
    add_paragraph(doc, "Для демонстрации полноценного чтения внутри мобильного приложения подключена открытая библиотека Project Gutenberg. PocketLib хранит стартовую подборку книг с обложками и идентификаторами, а текст загружает по запросу пользователя и отображает во встроенном ридере. В приложении также предусмотрены роли студента, преподавателя и администратора, локализация интерфейса, управление пользователями, добавление материалов и отдельный раздел официальных ссылок.")
    add_paragraph(doc, "Ключевые слова: электронная библиотека, мобильное приложение, React Native, Expo, TypeScript, SQLite, Node.js, XML, Project Gutenberg, Gutendex, API, колледж.")
    doc.add_page_break()


def contents(doc):
    doc.add_heading("СОДЕРЖАНИЕ", level=1)
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0)
    add_toc(p)
    doc.add_page_break()


def introduction(doc):
    doc.add_heading("ВВЕДЕНИЕ", level=1)
    add_paragraph(doc, "Цифровизация библиотечного фонда является актуальной задачей для образовательных организаций. Даже при наличии автоматизированной библиотечной информационной системы часть работы читателя остается привязанной к локальному компьютеру: поиск выполняется через отдельную программу, электронные материалы хранятся в папках, а доступ к открытым источникам не объединен в одном интерфейсе. Для студента удобнее иметь мобильное приложение, в котором можно открыть каталог, найти книгу, увидеть обложку и продолжить чтение с сохраненной страницы.")
    add_paragraph(doc, "Исходной точкой проекта стала резервная копия программы «Библиотечное дело» компании KostanaySoft. В ней обнаружены XML-файлы «КНИГИ_XML.XML» и «РЕЕСТР_XML.XML», а также локальные каталоги программы. Анализ показал, что XML содержит библиографические записи: название, авторов, издательство, год, ISBN, инвентарные сведения и стоимость. Полных текстов большинства книг в резервной копии нет. Поэтому проект PocketLib решает две разные задачи: сохраняет каталог существующего фонда и отдельно подключает только те электронные версии, использование которых допустимо.")
    add_paragraph(doc, "Практическая часть состоит из мобильного приложения PocketLib и локального сервиса BK Library API. Приложение ориентировано на Android и разрабатывается на React Native с использованием Expo. Оно поддерживает библиотеку, поиск, карточку книги, встроенный постраничный ридер, закладки, прогресс, локальные файлы, роли пользователей и административную панель. API работает с XML-выгрузкой, официальными источниками Казахстана и кэшированными результатами.")
    add_paragraph(doc, "Для наполнения демонстрационной библиотеки полными текстами выбран Project Gutenberg. Это позволяет показать чтение непосредственно в PocketLib, не прибегая к сомнительным загрузчикам и обходу защиты внешних сайтов. Важным ограничением проекта является соблюдение условий конкретного источника: для производственного использования необходимо проверять применимое авторское право, учитывать правила доступа Project Gutenberg и при массовой работе размещать разрешенные копии или собственный каталог на сервере организации.")
    add_paragraph(doc, "Цель дипломного проекта — разработать мобильное приложение электронной библиотеки PocketLib, объединяющее каталог учебной организации, открытые электронные материалы, встроенное чтение и административное управление.")
    add_paragraph(doc, "Для достижения цели поставлены следующие задачи:")
    add_bullets(doc, [
        "проанализировать структуру резервной копии АБИС и определить состав доступных метаданных;",
        "спроектировать архитектуру мобильного приложения и локального API;",
        "создать локальную базу данных для книг, пользователей и справочников;",
        "реализовать авторизацию и разграничение возможностей по ролям;",
        "разработать каталог, поиск, карточку книги и встроенный ридер;",
        "добавить легальный источник демонстрационных полных текстов и обложек;",
        "создать модуль работы с официальными ссылками на электронные учебники Казахстана;",
        "проверить сборку Android-приложения и подготовить документацию проекта.",
    ])


def general_part(doc):
    add_main_heading(doc, "1 ОБЩАЯ ЧАСТЬ")
    doc.add_heading("1.1 Обоснование необходимости разработки", level=2)
    add_paragraph(doc, "Библиотека колледжа обслуживает студентов, преподавателей и сотрудников. Ее фонд включает учебники, пособия, художественную литературу и методические материалы. Традиционная автоматизация хорошо решает задачу учета экземпляров, но не всегда предоставляет удобный мобильный доступ. Пользователь вынужден обращаться к библиотекарю, запускать настольное приложение или вручную искать электронную версию в интернете.")
    add_paragraph(doc, "В резервной копии исходной АБИС хранится каталог, сформированный за годы эксплуатации. Он ценен как база метаданных, однако сам по себе не превращается в онлайн-библиотеку. Наличие ISBN-папки или XML-записи не означает наличие электронного файла. Следовательно, новая система должна различать библиографическую карточку, внешнюю официальную ссылку, загруженный локальный файл и книгу с доступным открытым текстом.")
    add_paragraph(doc, "PocketLib создается как мобильная точка доступа. Пользователь получает единый интерфейс для навигации по материалам и чтения. Администратор управляет учетными записями и может пополнять локальную библиотеку. Преподаватель добавляет учебные материалы и распределяет их по дисциплинам. Отдельный раздел показывает официальные электронные ссылки, не смешивая их с книгами, которые доступны внутри приложения.")

    doc.add_heading("1.2 Цели и задачи системы", level=2)
    add_paragraph(doc, "Функциональная цель PocketLib заключается в сокращении времени от поиска материала до чтения. Система должна быть понятной для студента, достаточно гибкой для преподавателя и управляемой для администратора.")
    add_table(doc, ["№", "Задача", "Результат"], [
        ("1", "Импортировать существующий каталог", "Метаданные фонда доступны через BK Library API"),
        ("2", "Организовать мобильную библиотеку", "Каталог книг хранится в SQLite и отображается в приложении"),
        ("3", "Добавить внутреннее чтение", "TXT-материалы и книги Gutenberg открываются во встроенном ридере"),
        ("4", "Сохранить прогресс", "Страница, шрифт и закладки сохраняются локально"),
        ("5", "Разделить права", "Роли student, teacher и admin ограничивают доступ к операциям"),
        ("6", "Подключить официальные ссылки", "Отдельный экран работает с фильтром BK API"),
        ("7", "Проверить качество сборки", "TypeScript и Android bundle проходят проверку"),
    ], [1, 7, 8])

    doc.add_heading("1.3 Анализ исходных данных", level=2)
    add_paragraph(doc, "В папке резервной копии присутствуют XML-файлы «КНИГИ_XML.XML», «РЕЕСТР_XML.XML» и справочники. Основные файлы содержат каталог книг и реестр экземпляров. В ходе запуска API загружается 7010 карточек книг и 2101 запись реестра, суммарно 9111 элементов каталога. Эти числа описывают записи фонда, а не число готовых электронных файлов.")
    add_paragraph(doc, "Папка ISBN предназначена для локальных электронных материалов старой программы. Если внутри ISBN-папки существуют index.htm, обложки и страницы, материал можно открывать через локальный сервер. Однако резервная копия не содержит полные страницы для большинства записей. Поэтому попытка автоматически показать содержание каждой XML-карточки приводит к ошибке отсутствующего файла.")
    add_table(doc, ["Источник", "Назначение", "Содержимое", "Использование в PocketLib"], [
        ("КНИГИ_XML.XML", "Каталог фонда", "Название, автор, издательство, год, ISBN, инвентарные поля", "Импорт метаданных"),
        ("РЕЕСТР_XML.XML", "Реестр записей", "Учетные и библиографические сведения", "Дополнительный источник каталога"),
        ("ISBN/{ISBN}", "Локальные электронные папки", "index.htm, изображения страниц, обложки, HTML и ресурсы", "Открытие только реально существующих файлов"),
        ("gov.kz / data.egov.kz", "Официальные ссылки", "Перечни и ссылки на электронные учебники", "Отдельный фильтр официального доступа"),
        ("Project Gutenberg", "Открытая литература", "Тексты и обложки произведений", "Демонстрационная внутренняя библиотека"),
    ], [3, 4, 6, 5])

    doc.add_heading("1.4 Требования к системе", level=2)
    add_paragraph(doc, "К основным функциональным требованиям относятся:")
    add_bullets(doc, [
        "регистрация и вход пользователя;",
        "просмотр библиотеки, поиск и фильтрация материалов;",
        "отображение названия, автора, обложки, источника и описания;",
        "встроенное чтение текстовых книг с разбиением на страницы;",
        "сохранение закладок и позиции чтения;",
        "ручное добавление файлов преподавателем или администратором;",
        "административное управление пользователями и ролями;",
        "подключение стартовой подборки Gutenberg с обложками;",
        "просмотр официальных ссылок Казахстана отдельным разделом;",
        "работа приложения без обязательного запуска BK API для основной полки Gutenberg.",
    ])
    add_paragraph(doc, "Нефункциональные требования включают понятный мобильный интерфейс, сохранение данных между перезапусками, отсутствие пиратского загрузчика, отказ от обхода flipbook/viewer-защиты, ограничение опасных загрузок, возможность расширения каталога и воспроизводимую Android-сборку.")

    doc.add_heading("1.5 Выбор технологий", level=2)
    add_table(doc, ["Технология", "Назначение", "Причина выбора"], [
        ("React Native 0.81", "Мобильный интерфейс", "Единая кодовая база и нативные компоненты"),
        ("Expo SDK 54", "Среда разработки и сборки", "Быстрый запуск, готовые модули устройства"),
        ("TypeScript 5.9", "Язык клиентской части", "Статическая типизация и снижение числа ошибок"),
        ("Expo Router", "Навигация", "Файловая структура маршрутов"),
        ("React Native Paper", "UI-компоненты", "Готовые элементы Material Design"),
        ("expo-sqlite", "Локальная БД", "Сохранение каталога и справочников между запусками"),
        ("AsyncStorage", "Настройки чтения", "Простое хранение прогресса и закладок"),
        ("Node.js HTTP API", "Локальный сервер", "Импорт XML и выдача REST-подобных endpoint"),
        ("Project Gutenberg + Gutendex", "Открытая литература", "Метаданные, обложки и доступные тексты"),
    ], [4, 5, 8])


def architecture_part(doc):
    add_main_heading(doc, "2 ПРОЕКТИРОВАНИЕ СИСТЕМЫ")
    doc.add_heading("2.1 Общая архитектура", level=2)
    add_paragraph(doc, "PocketLib использует комбинированную архитектуру. Основная мобильная библиотека способна работать автономно после заполнения локальной SQLite-базы. Стартовый каталог Gutenberg добавляется непосредственно приложением. BK Library API подключается отдельно: он нужен для просмотра XML-фонда и официальных ссылок, но не блокирует чтение открытой литературы.")
    doc.add_picture(str(ARCHITECTURE_IMAGE), width=Cm(16))
    add_caption(doc, "Рисунок 1 — Архитектура PocketLib")
    add_paragraph(doc, "Такое разделение решает проблему неполных исходных данных. Библиографические карточки фонда не выдаются за электронные книги. Материалы с полным текстом отображаются как доступные для чтения внутри приложения. Официальные ссылки издательств остаются в отдельном онлайн-разделе.")

    doc.add_heading("2.2 Ролевая модель", level=2)
    add_paragraph(doc, "В системе реализовано три роли. При регистрации новая учетная запись получает роль студента. Повышение прав выполняется администратором из специальной панели. Это исключает самостоятельное назначение привилегий.")
    doc.add_picture(str(ROLES_IMAGE), width=Cm(16))
    add_caption(doc, "Рисунок 2 — Ролевая модель PocketLib")
    add_table(doc, ["Операция", "Студент", "Преподаватель", "Администратор"], [
        ("Просмотр каталога", "Да", "Да", "Да"),
        ("Чтение книг", "Да", "Да", "Да"),
        ("Закладки и прогресс", "Да", "Да", "Да"),
        ("Добавление файла", "Нет", "Да", "Да"),
        ("Назначение дисциплины", "Нет", "Да", "Да"),
        ("Просмотр списка пользователей", "Нет", "Нет", "Да"),
        ("Изменение ролей", "Нет", "Нет", "Да"),
        ("Удаление учетных записей", "Нет", "Нет", "Да"),
        ("Синхронизация gov.kz", "Нет", "Нет", "Да"),
    ], [7, 3, 4, 4])

    doc.add_heading("2.3 Локальная база данных", level=2)
    add_paragraph(doc, "Локальная база открывается через expo-sqlite и хранится в файле pocketlib.db. Инициализация выполняется при запуске приложения. Если таблиц еще нет, они создаются автоматически. Дополнительные поля книг добавляются миграциями ALTER TABLE, поэтому ранее установленное приложение может обновляться без удаления пользовательских данных.")
    doc.add_picture(str(DATABASE_IMAGE), width=Cm(16))
    add_caption(doc, "Рисунок 3 — Логическая схема базы данных PocketLib")
    add_table(doc, ["Таблица", "Назначение", "Ключевые поля"], [
        ("users", "Учетные записи и роли", "full_name, email, password, role, group_name"),
        ("books", "Локальный каталог", "title, author, cover_url, file_path, source, gutenberg_id"),
        ("disciplines", "Учебные дисциплины", "name, color"),
        ("courses", "Курсы дисциплин", "name, year, discipline_id"),
        ("specialities", "Специальности", "name"),
        ("categories", "Категории материалов", "name"),
        ("reading_history", "История открытия", "user_id, book_id, progress"),
        ("search_cache", "Кэш поиска", "query, results_json, timestamp"),
    ], [3, 5, 9])

    doc.add_heading("2.4 Навигация мобильного приложения", level=2)
    add_paragraph(doc, "Expo Router формирует маршруты на основе файлов в каталоге app. После входа пользователь попадает в набор вкладок. Доступность отдельных вкладок зависит от роли.")
    add_table(doc, ["Маршрут", "Экран", "Назначение"], [
        ("/(tabs)/index", "Главная", "Статистика, быстрые действия и последние книги"),
        ("/(tabs)/library", "Библиотека", "Каталог, фильтры и обновление Gutenberg"),
        ("/(tabs)/official", "Официальные", "Ссылки gov.kz, data.egov.kz и кандидаты"),
        ("/(tabs)/gutendex", "Мировая база", "Поиск открытой литературы Gutendex"),
        ("/(tabs)/add", "Добавление", "Импорт локального материала преподавателем или администратором"),
        ("/(tabs)/admin", "Администратор", "Пользователи, роли и синхронизация"),
        ("/book/[id]", "Карточка книги", "Метаданные, обложка и переход к чтению"),
        ("/reader/[id]", "Ридер", "Постраничное чтение и закладки"),
    ], [4, 4, 9])

    doc.add_heading("2.5 Сценарии использования", level=2)
    add_numbered(doc, [
        "Студент входит в приложение, открывает библиотеку, выбирает книгу Gutenberg, нажимает «Читать в приложении», меняет размер шрифта и добавляет закладку.",
        "Преподаватель открывает экран добавления, выбирает TXT или PDF-файл с устройства, заполняет метаданные и сохраняет материал в локальной библиотеке.",
        "Администратор входит под учетной записью администратора, открывает панель пользователей, назначает роль преподавателя и при необходимости удаляет обычную учетную запись.",
        "Администратор запускает локальный BK API и открывает вкладку официальных ссылок, чтобы получить витрину учебников и выполнить классификацию.",
        "Пользователь вводит английское название или автора во вкладке Gutendex и добавляет найденную книгу в личную локальную библиотеку.",
    ])


def implementation_part(doc):
    add_main_heading(doc, "3 СПЕЦИАЛЬНАЯ ЧАСТЬ")
    doc.add_heading("3.1 Структура мобильного проекта", level=2)
    add_paragraph(doc, "Проект pocketlib-rn организован по модульному принципу. Каталог app содержит экраны и маршруты, components — повторно используемые карточки и бейджи, services — доступ к данным и внешним источникам, contexts — состояние авторизации и языка, constants — тему интерфейса.")
    add_code(doc, """
pocketlib-rn/
  app/
    (tabs)/        # основные вкладки
    auth/          # вход и регистрация
    book/[id].tsx  # карточка книги
    reader/[id].tsx
  components/      # карточки и бейджи
  contexts/        # AuthContext, LanguageContext
  services/        # SQLite, книги, API, ридер, пользователи
  assets/          # иконки приложения
""")

    doc.add_heading("3.2 Авторизация и пользователи", level=2)
    add_paragraph(doc, "AuthContext хранит текущего пользователя и предоставляет операции входа, регистрации и выхода. Сервис userService обращается к таблице users. Пароль в текущем локальном прототипе сохраняется в SQLite. Для производственного развертывания необходимо перенести аутентификацию на сервер, хранить только криптографические хэши и добавить безопасное восстановление доступа.")
    add_paragraph(doc, "При первом запуске создается локальная административная учетная запись. Она предназначена для демонстрации и первичной настройки. Перед реальным использованием пароль необходимо заменить, а встроенные демонстрационные реквизиты удалить из кода.")
    add_table(doc, ["Поле пользователя", "Назначение"], [
        ("full_name", "Отображаемое имя"),
        ("email", "Уникальный логин"),
        ("password", "Пароль прототипа; требует хэширования в production"),
        ("role", "student, teacher или admin"),
        ("speciality_id", "Связь со специальностью"),
        ("course_id", "Связь с курсом"),
        ("group_name", "Учебная группа"),
    ], [5, 11])

    doc.add_heading("3.3 Каталог книг", level=2)
    add_paragraph(doc, "Сервис bookService инкапсулирует операции SQLite. Метод getAllBooks формирует запрос с фильтрами по дисциплине, курсу, категории, специальности, типу материала, языку, семестру, строке поиска и признаку офлайн-доступа. Метод addBook используется для ручного импорта и для сохранения книг из открытых источников.")
    add_paragraph(doc, "Карточка книги отображает обложку через expo-image, название, автора, бейджи, источник, описание и информацию о возможности внутреннего чтения. Если заполнено поле file_path или ol_key, кнопка переводит пользователя во встроенный ридер. Если существует только external_url, приложение открывает официальный внешний адрес.")
    add_code(doc, """
const internalReaderAvailable = Boolean(book.file_path || book.ol_key);

if (book.file_path || book.ol_key) {
  router.push(`/reader/${book.id}`);
} else if (book.external_url) {
  await Linking.openURL(book.external_url);
}
""")

    doc.add_heading("3.4 Подборка Gutenberg и поиск Gutendex", level=2)
    add_paragraph(doc, "Основная демонстрационная полка PocketLib заполняется открытыми книгами Project Gutenberg. В services/api.ts хранится небольшая стартовая подборка известных произведений. Для каждой записи задаются идентификатор Gutenberg, название, автор и URL обложки. Такой подход позволяет сразу показать полноценный сценарий приложения, не выдавая библиографическую XML-карточку за доступный электронный файл.")
    add_paragraph(doc, "Метод syncGutenbergBooks добавляет книги в SQLite и обновляет существующие записи по gutenberg_id. Метод searchBooks обращается к Gutendex, который возвращает метаданные и форматы Project Gutenberg. При медленном соединении поиск использует локальную стартовую подборку как fallback.")
    add_paragraph(doc, "В текущем прототипе текст загружается при открытии книги. Для производственного внедрения необходимо учитывать правила Project Gutenberg: при массовом использовании не следует создавать нагрузку прямыми глубокими ссылками. Правильный следующий шаг — загрузить разрешенный каталог и допустимые копии на собственный сервер колледжа либо развернуть собственный экземпляр Gutendex и контролируемый файловый кэш.")
    add_table(doc, ["ID Gutenberg", "Произведение", "Автор"], [
        ("84", "Frankenstein; Or, The Modern Prometheus", "Mary Wollstonecraft Shelley"),
        ("1342", "Pride and Prejudice", "Jane Austen"),
        ("11", "Alice's Adventures in Wonderland", "Lewis Carroll"),
        ("1661", "The Adventures of Sherlock Holmes", "Arthur Conan Doyle"),
        ("345", "Dracula", "Bram Stoker"),
        ("2554", "Crime and Punishment", "Fyodor Dostoyevsky"),
        ("1260", "Jane Eyre: An Autobiography", "Charlotte Bronte"),
    ], [3, 8, 6])

    doc.add_heading("3.5 Встроенный ридер", level=2)
    add_paragraph(doc, "ReaderScreen открывает локальный TXT-файл либо загружает текст Gutenberg по идентификатору. Перед отображением выполняется очистка строк и разбиение текста на страницы приблизительно по 1200 символов. Разбиение по абзацам снижает нагрузку на интерфейс и делает навигацию предсказуемой.")
    add_paragraph(doc, "Пользователь может перейти вперед и назад, увеличить или уменьшить шрифт, переключить тему чтения, добавить закладку и открыть список закладок. ReaderService сохраняет номер страницы, общее число страниц, размер шрифта и массив закладок в AsyncStorage. Ключ прогресса зависит от локального файла, URL или идентификатора Gutenberg.")
    add_table(doc, ["Функция", "Реализация"], [
        ("Пагинация", "Разбиение текста по абзацам с лимитом около 1200 символов"),
        ("Размер шрифта", "Диапазон 12–28 pt с шагом 2"),
        ("Темы", "light, sepia, dark"),
        ("Прогресс", "AsyncStorage по ключу книги"),
        ("Закладки", "Массив индексов страниц"),
        ("PDF", "Открытие через системный просмотрщик Android"),
        ("TXT", "Чтение внутри приложения"),
    ], [5, 12])

    doc.add_heading("3.6 Локализация интерфейса", level=2)
    add_paragraph(doc, "LanguageContext хранит выбранный язык интерфейса в AsyncStorage. Поддерживаются значения ru, kk и en. Переведены ключевые элементы навигации и административного интерфейса. Для полного завершения локализации следует перенести оставшиеся экранные строки в единый словарь и исключить литералы из компонентов.")
    add_table(doc, ["Код", "Язык", "Назначение"], [
        ("ru", "Русский", "Базовый язык интерфейса"),
        ("kk", "Казахский", "Локализация для пользователей колледжа"),
        ("en", "Английский", "Дополнительный язык интерфейса"),
    ], [3, 5, 9])

    doc.add_heading("3.7 BK Library API", level=2)
    add_paragraph(doc, "BK Library API — локальный Node.js HTTP-сервис на порту 3047. Он читает XML-выгрузки старой АБИС, предоставляет браузерный UI и JSON-endpoint. Сервер не является обязательным условием чтения Gutenberg-книг: приложение использует его только для XML-каталога и официального онлайн-раздела.")
    add_paragraph(doc, "Сервис обогащения не должен скачивать материалы с сомнительных сайтов или обходить защиту viewer. Для официальных источников сохраняется ссылка, тип доступа и степень уверенности совпадения. Прямая загрузка допускается только для разрешенных расширений и официальных URL без авторизации и без обхода защиты.")
    add_table(doc, ["Метод", "Endpoint", "Назначение"], [
        ("GET", "/api/health", "Проверка доступности сервера"),
        ("GET", "/api/books", "Поиск и пагинация каталога"),
        ("GET", "/api/xml-files", "Список XML-файлов"),
        ("GET", "/api/isbn-folders", "Список ISBN-папок"),
        ("POST", "/api/official/filter", "Формирование витрины официальных ссылок"),
        ("GET", "/api/official/stats", "Статистика официального фильтра"),
        ("GET", "/api/official/books", "Официально доступные книги"),
        ("GET", "/api/official/candidates", "Кандидаты для ручной проверки"),
        ("POST", "/api/official/classify", "Классификация по дисциплинам"),
        ("GET", "/api/books/{id}/content", "Получение локального файла или статуса"),
        ("POST", "/api/books/{id}/upload-content", "Ручная загрузка разрешенного файла"),
    ], [2, 7, 8])

    doc.add_heading("3.8 Фильтр официального доступа", level=2)
    add_paragraph(doc, "В officialAccessFilterService реализован сценарий, при котором система не пытается искать содержание для всех 9111 записей фонда. Вместо этого она формирует отдельную витрину до 100 официальных материалов. Источником служит публичный документ gov.kz и локальные кэши data/official-sources. Подход снижает сетевую нагрузку и риск ошибочного прикрепления сторонних материалов.")
    add_paragraph(doc, "Сопоставление использует нормализованные название, авторов, издательство, год и класс. Высокая уверенность позволяет прикрепить ссылку автоматически. Средняя уверенность переводит ссылку в candidateLinks для проверки администратором. Flipbook и онлайн-ридер сохраняются только как внешние URL.")
    add_table(doc, ["Статус", "Смысл"], [
        ("not_checked", "Запись еще не проверена"),
        ("not_available_online", "Официальная электронная версия не найдена"),
        ("official_link_found", "Найдена официальная ссылка"),
        ("online_reader_found", "Доступно чтение на сайте источника"),
        ("flipbook_found", "Найден flipbook; обход защиты запрещен"),
        ("downloadable_file_found", "Есть официальная прямая ссылка на файл"),
        ("manual_uploaded", "Файл добавлен администратором вручную"),
        ("candidate_found", "Требуется ручное подтверждение"),
        ("restricted_access", "Доступ ограничен правилами источника"),
    ], [6, 11])

    doc.add_heading("3.9 Безопасность", level=2)
    add_paragraph(doc, "Безопасность проекта рассматривается на уровне прототипа и будущего внедрения. Мобильное приложение не должно выполнять произвольный код из загруженных материалов. API ограничивает список допустимых расширений: PDF, EPUB, TXT и DOCX. Опасные типы EXE, JS, BAT, CMD и SH запрещены. Имя файла очищается, а путь сохранения формируется только внутри content/books/{bookId}. Максимальный размер загрузки — 100 МБ.")
    add_paragraph(doc, "Для производственной версии требуется усилить контур: вынести учетные записи на сервер, хэшировать пароли алгоритмом Argon2 или bcrypt, использовать HTTPS, добавить токены доступа, журналирование действий администратора, резервное копирование SQLite или переход на серверную СУБД, а также проверку MIME-типа и антивирусное сканирование загружаемых файлов.")
    add_table(doc, ["Риск", "Текущее ограничение", "Рекомендация для production"], [
        ("Утечка паролей", "Локальный прототип", "Серверная аутентификация и хэширование"),
        ("Опасный файл", "Белый список расширений", "MIME-проверка и антивирус"),
        ("Потеря каталога", "SQLite на устройстве", "Резервное копирование и серверная синхронизация"),
        ("Нарушение лицензии", "Только открытые и официальные источники", "Юридическая проверка и собственный файловый кэш"),
        ("Чрезмерная нагрузка на источник", "Стартовая подборка и запрос по открытию", "Собственный Gutendex и зеркало разрешенных файлов"),
    ], [5, 6, 7])

    doc.add_heading("3.10 Запуск и сборка", level=2)
    add_paragraph(doc, "Для разработки мобильного приложения используется Expo. Основная библиотека Gutenberg запускается без BK API. Локальный сервер нужен только при работе с XML и вкладкой официальных источников.")
    add_code(doc, """
# Мобильное приложение
cd pocketlib-rn
npm install
npm start

# Дополнительный локальный API
cd ..\\bk-api
npm start

# Проверка типов
cd ..\\pocketlib-rn
npx tsc --noEmit

# Проверка Android bundle
npx expo export --platform android --output-dir .expo-export-check-android
""")

    doc.add_heading("3.11 Тестирование", level=2)
    add_paragraph(doc, "В процессе разработки выполнена статическая проверка TypeScript и тестовая экспортная сборка Android bundle. Дополнительно проверены освобождение порта 3047 после остановки API, наличие обложек Gutenberg и формирование маршрута к встроенному ридеру.")
    add_table(doc, ["№", "Проверка", "Ожидаемый результат", "Результат"], [
        ("1", "npx tsc --noEmit", "Нет ошибок TypeScript", "Пройдено"),
        ("2", "expo export --platform android", "Android bundle формируется", "Пройдено"),
        ("3", "Открытие карточки Gutenberg", "Есть обложка и кнопка внутреннего чтения", "Пройдено по коду и сборке"),
        ("4", "Остановка BK API", "Порт 3047 освобожден", "Пройдено"),
        ("5", "Ридер Gutenberg", "Текст разбит на страницы, доступны закладки", "Реализовано"),
        ("6", "Роли", "Студент не видит административную вкладку", "Реализовано"),
        ("7", "Официальные ссылки", "Открываются отдельным экраном", "Реализовано"),
    ], [1, 5, 6, 5])
    add_paragraph(doc, "Для приемочного тестирования на физическом Android-устройстве рекомендуется дополнительно проверить работу в нестабильной сети, повторное открытие книги после перезапуска приложения, импорт локального TXT/PDF, локализацию, изменение ролей и поведение при недоступном BK API.")


def economic_part(doc):
    add_main_heading(doc, "4 ЭКОНОМИЧЕСКАЯ ЧАСТЬ")
    doc.add_heading("4.1 Оценка трудозатрат", level=2)
    add_paragraph(doc, "Экономическая оценка приведена как пример для дипломного проекта. Перед сдачей значения следует согласовать с методическими требованиями колледжа и при необходимости заменить тарифы актуальными данными организации.")
    rows = [
        ("Анализ исходной АБИС и XML", "36"),
        ("Проектирование базы и архитектуры", "42"),
        ("Разработка BK Library API", "64"),
        ("Разработка экранов мобильного приложения", "110"),
        ("Реализация ридера и Gutenberg", "58"),
        ("Административная панель и роли", "42"),
        ("Тестирование и исправления", "48"),
        ("Документация", "40"),
    ]
    add_table(doc, ["Этап", "Часы"], rows, [12, 4])
    total_hours = sum(int(row[1]) for row in rows)
    add_paragraph(doc, f"Общая расчетная трудоемкость проекта составляет {total_hours} часов. При условной ставке 2000 тенге за час стоимость разработки равна {total_hours} × 2000 = {total_hours * 2000:,} тенге.".replace(",", " "))

    doc.add_heading("4.2 Расчет прямых затрат", level=2)
    add_table(doc, ["Статья", "Расчет", "Сумма, тг"], [
        ("Оплата труда разработчика", f"{total_hours} ч × 2000 тг", f"{total_hours * 2000:,}".replace(",", " ")),
        ("Электроэнергия", f"0,25 кВт × {total_hours} ч × 30 тг", f"{int(0.25 * total_hours * 30):,}".replace(",", " ")),
        ("Интернет", "2 месяца × 8000 тг", "16 000"),
        ("Амортизация оборудования", "условная оценка", "45 000"),
        ("Прочие расходы", "расходные материалы и резерв", "20 000"),
    ], [7, 6, 4])
    total_cost = total_hours * 2000 + int(0.25 * total_hours * 30) + 16000 + 45000 + 20000
    add_paragraph(doc, f"Итоговая расчетная стоимость разработки прототипа составляет {total_cost:,} тенге.".replace(",", " "))

    doc.add_heading("4.3 Оценка эффекта внедрения", level=2)
    add_paragraph(doc, "PocketLib не заменяет библиотекаря и действующую АБИС. Экономический эффект возникает за счет сокращения повторяющихся действий: поиска электронной версии, ручной передачи файлов, объяснения доступа к материалам, проверки одинаковых ссылок и восстановления позиции чтения. Дополнительный эффект связан с использованием существующего XML-каталога вместо ручного ввода всех карточек.")
    add_table(doc, ["Операция", "До внедрения", "После внедрения", "Эффект"], [
        ("Поиск открытой книги", "Ручной поиск по сайтам", "Поиск и карточка в приложении", "Снижение времени"),
        ("Передача учебной ссылки", "Сообщения и списки", "Единый официальный раздел", "Меньше дублирования"),
        ("Чтение TXT", "Отдельный просмотрщик", "Встроенный ридер", "Удобство пользователя"),
        ("Возврат к странице", "Ручной поиск", "Сохраненный прогресс", "Экономия времени"),
        ("Каталог фонда", "Настольная АБИС", "JSON API и мобильный интерфейс", "Расширение доступности"),
    ], [5, 6, 6, 5])

    doc.add_heading("4.4 Возможности развития", level=2)
    add_paragraph(doc, "При переходе от прототипа к внедрению целесообразно развернуть серверную базу данных, безопасную аутентификацию, резервное копирование, собственный файловый кэш разрешенных материалов и панель библиотекаря. Это потребует дополнительных затрат, но позволит использовать PocketLib как внутренний сервис колледжа для нескольких групп и устройств.")


def safety_part(doc):
    add_main_heading(doc, "5 ОХРАНА ТРУДА И ТЕХНИКА БЕЗОПАСНОСТИ")
    doc.add_heading("5.1 Условия труда разработчика", level=2)
    add_paragraph(doc, "Работа над программным продуктом выполняется преимущественно за персональным компьютером. Основными факторами нагрузки являются длительное сидячее положение, зрительное напряжение, повторяющиеся движения кистей, статическая нагрузка на мышцы спины и шеи, а также эмоциональная нагрузка при поиске и исправлении ошибок.")
    add_paragraph(doc, "Рабочее место должно обеспечивать устойчивое положение корпуса, удобную высоту стола и кресла, достаточное пространство для ног, правильное расположение монитора и отсутствие бликов. Верхняя часть экрана располагается примерно на уровне глаз, а расстояние до монитора выбирается комфортным для чтения текста.")

    doc.add_heading("5.2 Освещение и микроклимат", level=2)
    add_paragraph(doc, "Помещение должно иметь равномерное освещение. Недопустимо размещать яркий источник света непосредственно за монитором или напротив него. При естественном освещении экран располагают так, чтобы окно не создавало отражений. Температура, влажность и вентиляция должны поддерживать комфортные условия длительной интеллектуальной работы.")

    doc.add_heading("5.3 Организация перерывов", level=2)
    add_paragraph(doc, "Во время разработки необходимо делать регулярные короткие перерывы, менять положение тела, выполнять упражнения для глаз и разминку кистей. Непрерывная работа без пауз снижает концентрацию и повышает вероятность ошибок в коде. Рациональный режим полезен не только для здоровья, но и для качества программного продукта.")

    doc.add_heading("5.4 Электробезопасность", level=2)
    add_paragraph(doc, "Компьютер, монитор, зарядные устройства и сетевое оборудование подключаются только к исправным розеткам. Кабели не должны иметь повреждений и располагаться в проходе. Нельзя самостоятельно вскрывать блоки питания под напряжением. При появлении запаха гари, искрения или сильного нагрева оборудование необходимо отключить.")

    doc.add_heading("5.5 Пожарная безопасность", level=2)
    add_paragraph(doc, "Перегрузка удлинителей, неисправные кабели, перекрытая вентиляция и накопление пыли увеличивают риск возгорания. Рабочее место необходимо содержать в порядке, не размещать бумагу и жидкости рядом с источниками питания. В помещении должен быть исправный огнетушитель, предназначенный для электрооборудования. Тушить включенную технику водой запрещено.")

    doc.add_heading("5.6 Экологические аспекты", level=2)
    add_paragraph(doc, "PocketLib способствует сокращению необязательной печати списков и копий материалов. Электронный каталог и встроенное чтение уменьшают расход бумаги и картриджей. Одновременно важно рационально использовать электроэнергию, выключать неиспользуемое оборудование и сдавать электронные отходы в специализированные пункты приема.")

    doc.add_heading("5.7 Информационная безопасность рабочего места", level=2)
    add_paragraph(doc, "На рабочем компьютере могут находиться исходный код, XML-каталог, резервные копии и пользовательские данные. Необходимо ограничивать доступ посторонних лиц, использовать пароль учетной записи, регулярно создавать резервные копии и перед утилизацией накопителя безопасно удалять информацию. При публикации демонстрационных материалов нельзя раскрывать персональные сведения читателей.")


def conclusion(doc):
    add_main_heading(doc, "ЗАКЛЮЧЕНИЕ")
    add_paragraph(doc, "В результате дипломного проекта разработан прототип мобильного приложения PocketLib для электронной библиотеки колледжа. Выполнен анализ резервной копии программы «Библиотечное дело» KostanaySoft и установлено, что XML-выгрузка содержит преимущественно метаданные фонда, а не полные тексты. Это определило архитектуру решения: каталог и электронный контент обрабатываются раздельно.")
    add_paragraph(doc, "Мобильное приложение реализовано на React Native, Expo и TypeScript. Созданы экраны входа, регистрации, главной панели, библиотеки, официальных ссылок, мировой базы Gutendex, добавления книг, профиля, карточки материала, встроенного ридера и панели администратора. Локальная SQLite-база хранит книги, пользователей и справочники. AsyncStorage сохраняет прогресс чтения, размер шрифта и закладки.")
    add_paragraph(doc, "Для демонстрации внутреннего чтения подключена стартовая подборка Project Gutenberg с обложками. Тексты открываются непосредственно в PocketLib, разбиваются на страницы и отображаются с возможностью менять тему и размер шрифта. Отдельно реализован BK Library API на Node.js: он предоставляет доступ к XML-каталогу и витрине официальных ссылок Казахстана, не пытаясь скачивать пиратские материалы и не обходя защиту внешних ридеров.")
    add_paragraph(doc, "Проверка TypeScript и экспорт Android bundle выполнены успешно. Проект можно развивать дальше: добавить серверную авторизацию, хэширование паролей, централизованную базу данных, серверный кэш разрешенных текстов, полноценную локализацию всех строк, журнал действий администратора и публикацию приложения для Android.")


def sources(doc):
    add_main_heading(doc, "СПИСОК ИСПОЛЬЗОВАННОЙ ЛИТЕРАТУРЫ И ИСТОЧНИКОВ")
    sources_list = [
        "React Native Documentation. Introduction. URL: https://reactnative.dev/docs/getting-started (дата обращения: 31.05.2026).",
        "Expo Documentation. Expo Router. URL: https://docs.expo.dev/router/introduction (дата обращения: 31.05.2026).",
        "Expo Documentation. SQLite. URL: https://docs.expo.dev/versions/latest/sdk/sqlite (дата обращения: 31.05.2026).",
        "Expo Documentation. FileSystem. URL: https://docs.expo.dev/versions/latest/sdk/filesystem (дата обращения: 31.05.2026).",
        "TypeScript Documentation. URL: https://www.typescriptlang.org/docs (дата обращения: 31.05.2026).",
        "Node.js Documentation. HTTP. URL: https://nodejs.org/api/http.html (дата обращения: 31.05.2026).",
        "Project Gutenberg. Free eBooks. URL: https://www.gutenberg.org (дата обращения: 31.05.2026).",
        "Project Gutenberg. Terms of Use. URL: https://www.gutenberg.org/policy/terms_of_use.html (дата обращения: 31.05.2026).",
        "Project Gutenberg. Collection Development. URL: https://www.gutenberg.org/policy/collection_development.html (дата обращения: 31.05.2026).",
        "Gutendex: Web API for Project Gutenberg ebook metadata. URL: https://github.com/garethbjohnson/gutendex (дата обращения: 31.05.2026).",
        "Публичный документ Министерства просвещения Республики Казахстан. URL: https://www.gov.kz/memleket/entities/edu/documents/details/892700?lang=ru (дата обращения: 31.05.2026).",
        "Документация проекта PocketLib: исходный код pocketlib-rn и bk-api.",
        "XML-выгрузка АБИС «Библиотечное дело» KostanaySoft: КНИГИ_XML.XML, РЕЕСТР_XML.XML и справочники.",
        "Материалы по проектированию локальных баз данных SQLite и мобильных информационных систем.",
        "Методические материалы по охране труда при работе с персональным компьютером.",
    ]
    for idx, source in enumerate(sources_list, start=1):
        add_paragraph(doc, f"{idx}. {source}")


def appendices(doc):
    add_main_heading(doc, "ПРИЛОЖЕНИЕ А. ОСНОВНЫЕ ENDPOINT BK LIBRARY API")
    add_code(doc, """
GET  /api/health
GET  /api/books?source=all&q=физика&limit=20&offset=0
GET  /api/books/books/{id}
GET  /api/books/registry/{id}
GET  /api/xml-files
GET  /api/isbn-folders

POST /api/official/filter
GET  /api/official/stats
GET  /api/official/books
GET  /api/official/candidates
POST /api/official/classify

GET  /api/books/{id}/content
POST /api/books/{id}/upload-content
POST /api/books/{id}/attach-candidate
POST /api/books/{id}/reject-candidate
""")
    add_paragraph(doc, "Endpoint предназначены для локальной разработки и демонстрации. Для публикации API следует добавить HTTPS, серверную авторизацию и журналирование операций.")

    add_main_heading(doc, "ПРИЛОЖЕНИЕ Б. ПАМЯТКА ПО ЗАПУСКУ")
    add_numbered(doc, [
        "Установить Node.js и зависимости проекта командой npm install в папке pocketlib-rn.",
        "Запустить Expo командой npm start.",
        "Открыть проект на Android-устройстве через Expo Go или собрать приложение.",
        "Для основной Gutenberg-полки BK API не требуется.",
        "Для XML-каталога и официального онлайн-раздела запустить npm start в папке bk-api.",
        "При тестировании с телефона указать EXPO_PUBLIC_BK_API_URL с IP-адресом компьютера в одной Wi-Fi-сети.",
    ])
    add_code(doc, """
$env:EXPO_PUBLIC_BK_API_URL="http://192.168.8.61:3047"
cd C:\\Users\\Acer\\OneDrive\\Документы\\bk\\bk-api
npm start
""")

    add_main_heading(doc, "ПРИЛОЖЕНИЕ В. ОГРАНИЧЕНИЯ ПРОТОТИПА")
    add_bullets(doc, [
        "XML-фонд не содержит полные тексты большинства книг.",
        "Часть официальных учебников доступна только через внешние сайты издательств.",
        "Project Gutenberg требует проверки применимого авторского права за пределами США.",
        "Для массового производственного использования Gutenberg нужен собственный кэш или зеркало согласно правилам источника.",
        "Пароли локального прототипа необходимо заменить серверной безопасной аутентификацией.",
        "Локализация интерфейса требует завершения для всех экранных строк.",
        "PDF открывается системным просмотрщиком Android; встроенный PDF-компонент можно добавить отдельным этапом.",
    ])


def build_document():
    build_diagrams()
    doc = Document()
    configure_document(doc)
    properties = doc.core_properties
    properties.title = "Дипломная документация PocketLib"
    properties.subject = "Мобильное приложение электронной библиотеки"
    properties.author = "[Ф.И.О. студента]"
    properties.keywords = "PocketLib, React Native, Expo, SQLite, Node.js, Project Gutenberg"
    properties.comments = "Сформировано генератором docs/generate_pocketlib_diploma.py"

    title_page(doc)
    annotation(doc)
    contents(doc)
    introduction(doc)
    general_part(doc)
    architecture_part(doc)
    implementation_part(doc)
    economic_part(doc)
    safety_part(doc)
    conclusion(doc)
    sources(doc)
    appendices(doc)
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build_document()
